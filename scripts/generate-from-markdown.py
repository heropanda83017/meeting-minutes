#!/usr/bin/env python3
"""
Markdown → 公文格式Word文档生成脚本
使用方法：
  python3 generate-from-markdown.py <输入.md> [输出.docx]

流程：
  1. 读取Markdown源文件
  2. 按层级规则识别标题(H1/H2/H3/H4)/表格/引用/分隔线/正文
  3. 按会议记录排版规范重新生成文档
  4. 表格格式化为Word表格（表头黑体加粗、数据仿宋）
  5. 引用格式化为楷体加粗斜体（高亮结论）
  6. 分隔线转换为排版分隔
  7. 数字字体设为 Times New Roman

依赖：
  pip install python-docx lxml
"""
import sys, re, os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


def set_run_style(run, font_name, size_pt, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_para_format(p, indent=True, line_spacing=28, align=None):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if align:
        p.alignment = align
    if indent:
        pPr = p._element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), '200')


def add_title(doc, text, font="宋体", size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_style(run, font, size)
    spacer = doc.add_paragraph()
    set_para_format(spacer, indent=False, line_spacing=14)


def bold_prefix_text(text):
    """Detect bold prefix pattern like '会议时间：' at start of line"""
    m = re.match(r'^([^：]+：)', text)
    return m.group(1) if m else None


def add_body(doc, text):
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    set_para_format(p, indent=True)
    
    # Handle bold markers **text**
    # Split by **...** patterns
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_style(run, '仿宋', 16, bold=True)
        else:
            # Check for colon prefix for bold field labels
            bp = bold_prefix_text(part)
            if bp and len(bp) < 20:
                rest = part[len(bp):]
                run1 = p.add_run(bp)
                set_run_style(run1, '仿宋', 16, bold=True)
                if rest:
                    run2 = p.add_run(rest)
                    set_run_style(run2, '仿宋', 16)
            else:
                run = p.add_run(part)
                set_run_style(run, '仿宋', 16)


def add_blockquote(doc, text):
    clean = re.sub(r'^>\s*', '', text).strip()
    p = doc.add_paragraph()
    set_para_format(p, indent=True)
    run = p.add_run(clean)
    set_run_style(run, '楷体', 16, bold=True, italic=True)


def add_separator(doc):
    p = doc.add_paragraph()
    set_para_format(p, indent=False)
    run = p.add_run('—' * 20)
    set_run_style(run, '仿宋', 12)
    spacer = doc.add_paragraph()
    set_para_format(spacer, indent=False, line_spacing=8)


def parse_table_lines(lines):
    """Parse markdown table into list of lists"""
    if not lines:
        return None
    # First line is header
    header = [c.strip() for c in lines[0].split('|') if c.strip()]
    if not header:
        return None
    # Skip separator line (|---|)
    data_start = 1
    if len(lines) > 1 and re.match(r'^\|[-:| ]+\|$', lines[1].strip()):
        data_start = 2
    data = []
    for line in lines[data_start:]:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            data.append(cells)
    return [header] + data if data else None


def add_table(doc, table_data):
    """table_data: list of lists, first row = header"""
    if not table_data:
        return
    rows = len(table_data)
    cols = max(len(r) for r in table_data)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j >= cols:
                break
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            run = p.add_run(str(cell_text))
            if i == 0:  # header row
                set_run_style(run, '黑体', 12, bold=True)
            else:
                set_run_style(run, '仿宋', 12)
    
    spacer = doc.add_paragraph()
    set_para_format(spacer, indent=False, line_spacing=8)
    return table


def detect_md_header(line):
    """Detect markdown header level. Returns (level, text) or None"""
    m = re.match(r'^(#{1,5})\s+(.+)$', line.strip())
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        return (level, text)
    return None


def add_header(doc, level, text):
    """Add header at appropriate nesting level"""
    if text.startswith('**') and text.endswith('**'):
        text = text[2:-2]
    
    if level == 1:
        # document title
        add_title(doc, text)
    elif level == 2:
        p = doc.add_paragraph()
        set_para_format(p, indent=True)
        run = p.add_run(text)
        set_run_style(run, '黑体', 16, bold=True)
    elif level == 3:
        p = doc.add_paragraph()
        set_para_format(p, indent=True)
        run = p.add_run(text)
        set_run_style(run, '楷体', 16, bold=False)
    elif level >= 4:
        p = doc.add_paragraph()
        set_para_format(p, indent=True)
        run = p.add_run(text)
        set_run_style(run, '仿宋', 16, bold=True)


def set_number_fonts(doc):
    """Set Times New Roman for all digits in paragraphs AND table cells"""
    for p in doc.paragraphs:
        for run in list(p.runs):
            _set_run_numbers(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in list(p.runs):
                        _set_run_numbers(run)


def _set_run_numbers(run):
    """Inner helper to replace digits in one run with Times New Roman"""
    txt = run.text
    if not txt or not re.search(r'\d', txt):
        return
    parts = re.split(r'(\d+)', txt)
    if len(parts) <= 1:
        return
    
    orig_font = run.font.name
    orig_size = run.font.size
    orig_bold = run.font.bold
    rpr = run._element.find(qn('w:rPr'))
    east_asia = None
    if rpr is not None:
        rfe = rpr.find(qn('w:rFonts'))
        if rfe is not None:
            east_asia = rfe.get(qn('w:eastAsia'))
    
    parent = run._element.getparent()
    idx = list(parent).index(run._element)
    parent.remove(run._element)
    
    new_elems = []
    for part in parts:
        ne = etree.SubElement(parent, qn('w:r'))
        rp = etree.SubElement(ne, qn('w:rPr'))
        rf = etree.SubElement(rp, qn('w:rFonts'))
        if re.fullmatch(r'\d+', part):
            rf.set(qn('w:ascii'), 'Times New Roman')
            rf.set(qn('w:hAnsi'), 'Times New Roman')
            rf.set(qn('w:eastAsia'), 'Times New Roman')
        else:
            fn = orig_font or '仿宋'
            rf.set(qn('w:ascii'), fn)
            rf.set(qn('w:hAnsi'), fn)
            rf.set(qn('w:eastAsia'), east_asia or fn)
        if orig_size:
            sv = str(int(orig_size.pt * 2))
            etree.SubElement(rp, qn('w:sz')).set(qn('w:val'), sv)
            etree.SubElement(rp, qn('w:szCs')).set(qn('w:val'), sv)
        if orig_bold:
            etree.SubElement(rp, qn('w:b'))
        t = etree.SubElement(ne, qn('w:t'))
        t.text = part
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_elems.append(ne)
    
    for ne in new_elems:
        parent.remove(ne)
    for j, ne in enumerate(new_elems):
        parent.insert(idx + j, ne)


def is_table_line(line):
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and s.count('|') >= 2


def is_table_separator(line):
    return bool(re.match(r'^\|[-:| ]+\|$', line.strip()))


def convert_markdown_to_docx(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()
    lines = text.split('\n')
    
    doc = Document()
    
    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    
    in_table = False
    table_lines = []
    first_header_encountered = True
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip pure empty lines
        if not stripped:
            if in_table:
                # End of table
                tbl = parse_table_lines(table_lines)
                if tbl:
                    add_table(doc, tbl)
                in_table = False
                table_lines = []
            i += 1
            continue
        
        # Separator (--- or ___)
        if re.match(r'^-{3,}$', stripped) or re.match(r'^_{3,}$', stripped):
            if in_table:
                tbl = parse_table_lines(table_lines)
                if tbl:
                    add_table(doc, tbl)
                in_table = False
                table_lines = []
            add_separator(doc)
            i += 1
            continue
        
        # Blockquote
        if stripped.startswith('> '):
            if in_table:
                tbl = parse_table_lines(table_lines)
                if tbl:
                    add_table(doc, tbl)
                in_table = False
                table_lines = []
            add_blockquote(doc, stripped)
            i += 1
            continue
        
        # Markdown header
        h = detect_md_header(line)
        if h:
            if in_table:
                tbl = parse_table_lines(table_lines)
                if tbl:
                    add_table(doc, tbl)
                in_table = False
                table_lines = []
            
            level, htext = h
            if first_header_encountered and level == 1:
                add_title(doc, htext)
                first_header_encountered = False
            elif level == 1:
                # Treat as H1 even if more than one
                add_header(doc, 2, htext)  # demote to H2
            else:
                add_header(doc, level, htext)
            i += 1
            continue
        
        # Table line
        if is_table_line(line):
            table_lines.append(stripped)
            in_table = True
            i += 1
            continue
        
        # Regular text (not in table, not header, not separator, not blockquote)
        if in_table:
            tbl = parse_table_lines(table_lines)
            if tbl:
                add_table(doc, tbl)
            in_table = False
            table_lines = []
        
        add_body(doc, stripped)
        i += 1
    
    # Handle last pending table
    if in_table:
        tbl = parse_table_lines(table_lines)
        if tbl:
            add_table(doc, tbl)
    
    # Set all digits to Times New Roman
    set_number_fonts(doc)
    
    doc.save(output_path)
    
    # === 后处理：落款右对齐 ===
    # 识别文档末尾的落款行（公司名+日期），设为右对齐+正文同款字体
    signoff_indicators = ['公司', '集团', '局', '部', '年', '月', '日']
    signoff_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 匹配：以公司/集团/局/部结尾 或 包含年月日的行
        if any(text.endswith(ind) for ind in ['公司', '集团', '局', '部', '中心', '处']):
            if len(text) < 30:  # 短文本，大概率是落款
                signoff_paragraphs.append(para)
        elif '年' in text and '月' in text and '日' in text and len(text) < 30:
            signoff_paragraphs.append(para)
    
    if signoff_paragraphs:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for para in signoff_paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(16)
                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    doc.save(output_path)
    size = os.path.getsize(output_path)
    print(f'✅ 已保存至 {output_path} ({size} 字节)')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python3 generate-from-markdown.py <输入.md> [输出.docx]')
        sys.exit(1)
    input_file = sys.argv[1]
    if not os.path.exists(input_file):
        print(f'❌ 文件不存在: {input_file}')
        sys.exit(1)
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '.docx')
    convert_markdown_to_docx(input_file, output_file)
