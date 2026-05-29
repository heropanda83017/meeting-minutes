#!/usr/bin/env python3
"""
会议记录格式化生成脚本
使用方法：
  python3 generate-formatted-docx.py <输入.docx> [输出.docx]

流程：
  1. 读取原文档纯文本（清除所有格式）
  2. 按层级规则识别标题/正文
  3. 按会议记录排版规范重新生成
  4. 将数字字体设为 Times New Roman
"""
import sys, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree


def set_style(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_format(p, indent=True):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值模式
    pf.line_spacing = Pt(28)                          # 固定值28磅
    pf.space_before = Pt(0)                           # 段前0磅
    pf.space_after = Pt(0)                            # 段后0磅
    if indent:
        # 使用 XML 设置首行缩进2字符宽度
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), '200')  # 2字符宽度


def detect_level(text):
    """按 §0.3 层级识别规则判断段落类型"""
    if not text.strip():
        return 'empty'
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'h1'
    if re.match(r'^（[一二三四五六七八九十\d]+）', text):
        return 'h2'
    if re.match(r'^\d+[、.]', text):
        return 'h3'
    return 'body'


def format_docx(input_path, output_path):
    # ① 读取原文档纯文本
    orig = Document(input_path)
    paragraphs = [p.text for p in orig.paragraphs]

    # ② 创建新文档，设置页面
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ③ 逐段处理
    for text in paragraphs:
        if not text.strip():
            continue

        level = detect_level(text)

        if level == 'empty':
            continue
        elif text == paragraphs[0] and len(text) < 50:
            # 文档总标题（通常为第一段且较短）
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_style(run, '宋体', 22)  # 方正小标宋不可用时用宋体
            doc.add_paragraph()  # 空行
        elif level == 'h1':
            p = doc.add_paragraph()
            set_paragraph_format(p, indent=True)
            run = p.add_run(text)
            set_style(run, '黑体', 16, bold=True)
        elif level == 'h2':
            p = doc.add_paragraph()
            set_paragraph_format(p, indent=True)
            run = p.add_run(text)
            set_style(run, '楷体', 16, bold=False)
        elif level == 'h3':
            p = doc.add_paragraph()
            set_paragraph_format(p, indent=True)
            run = p.add_run(text)
            set_style(run, '仿宋', 16, bold=True)
        else:
            # 正文：检测是否有加粗标签前缀（如"会议时间："）
            bold_prefix_match = re.match(r'^([^：]+：)', text)
            p = doc.add_paragraph()
            set_paragraph_format(p, indent=True)
            if bold_prefix_match:
                prefix = bold_prefix_match.group(1)
                rest = text[len(prefix):]
                run = p.add_run(prefix)
                set_style(run, '仿宋', 16, bold=True)
                run = p.add_run(rest)
                set_style(run, '仿宋', 16)
            else:
                run = p.add_run(text)
                set_style(run, '仿宋', 16)

    # ④ 数字字体设为 Times New Roman（XML 级操作）
    for p in doc.paragraphs:
        for run in list(p.runs):
            txt = run.text
            if not txt or not re.search(r'\d', txt):
                continue
            parts = re.split(r'(\d+)', txt)
            if len(parts) <= 1:
                continue

            orig_font = run.font.name
            orig_size = run.font.size
            orig_bold = run.font.bold
            rpr_elem = run._element.find(qn('w:rPr'))
            east_asia = None
            if rpr_elem is not None:
                rfe = rpr_elem.find(qn('w:rFonts'))
                if rfe is not None:
                    east_asia = rfe.get(qn('w:eastAsia'))

            parent = run._element.getparent()
            idx = list(parent).index(run._element)
            parent.remove(run._element)

            new_elems = []
            for part in parts:
                ne = etree.SubElement(parent, qn('w:r'))
                rp = etree.SubElement(ne, qn('w:rPr'))
                rf = etree.SubElement(rp, qn('w:rFonts'))
                if re.fullmatch(r'\d+', part):
                    rf.set(qn('w:ascii'), 'Times New Roman')
                    rf.set(qn('w:hAnsi'), 'Times New Roman')
                    rf.set(qn('w:eastAsia'), 'Times New Roman')
                else:
                    fn = orig_font or '仿宋'
                    rf.set(qn('w:ascii'), fn)
                    rf.set(qn('w:hAnsi'), fn)
                    rf.set(qn('w:eastAsia'), east_asia or fn)
                if orig_size:
                    sv = str(int(orig_size.pt * 2))
                    etree.SubElement(rp, qn('w:sz')).set(qn('w:val'), sv)
                    etree.SubElement(rp, qn('w:szCs')).set(qn('w:val'), sv)
                if orig_bold:
                    etree.SubElement(rp, qn('w:b'))
                t = etree.SubElement(ne, qn('w:t'))
                t.text = part
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_elems.append(ne)

            for ne in new_elems:
                parent.remove(ne)
            for j, ne in enumerate(new_elems):
                parent.insert(idx + j, ne)

    # ⑤ 后处理：落款右对齐\n    signoff_paragraphs = []\n    for para in doc.paragraphs:\n        text = para.text.strip()\n        if any(text.endswith(ind) for ind in ['公司', '集团', '局', '部', '中心', '处']):\n            if len(text) < 30:\n                signoff_paragraphs.append(para)\n        elif '年' in text and '月' in text and '日' in text and len(text) < 30:\n            signoff_paragraphs.append(para)\n\n    for para in signoff_paragraphs:\n        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT\n        for run in para.runs:\n            run.font.size = Pt(16)\n            run.font.name = '仿宋'\n            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')\n\n    doc.save(output_path)\n    print(f'✅ 已保存至 {output_path}')\n    return output_path\n\n\nif __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python3 generate-formatted-docx.py <输入.docx> [输出.docx]')
        sys.exit(1)
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '-格式化版.docx')
    format_docx(input_file, output_file)
